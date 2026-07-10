"""
Generate BAREAI pre-panel presentation (Word) and script in Somali.
Output: pre-panel/BAREAI_PrePanel_Presentation.docx
         pre-panel/BAREAI_PrePanel_Script_Somali.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "pre-panel"
OUT_DIR.mkdir(exist_ok=True)

AUTHORS = "Yaasiin Mohamuud Abdullaahi, Naima Abdiaziz Said, Nasteha Mohamuud Mohamed, Najma Muhidiin Mohamed"
TITLE = "BAREAI — Automatic Classification of Crime-Related Text Reports Using NLP"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5


def centered(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_slides() -> Document:
    doc = Document()
    style_doc(doc)

    # SLIDE 1 — Title
    page_break(doc) if False else None
    for _ in range(3):
        doc.add_paragraph()
    centered(doc, "JAMHURIYA UNIVERSITY OF SCIENCE AND TECHNOLOGY (JUST)", 11)
    centered(doc, "Faculty of Computer & Information Technology", 11)
    doc.add_paragraph()
    centered(doc, "PRE-PANEL PRESENTATION", 14, bold=True)
    doc.add_paragraph()
    centered(doc, TITLE, 13, bold=True)
    centered(doc, "BAREAI — Bare AI Crime Intelligence Platform", 12, bold=True)
    doc.add_paragraph()
    centered(doc, AUTHORS, 11)
    centered(doc, "August 2026", 11, bold=True)
    page_break(doc)

    # SLIDE 2
    heading(doc, "SLIDE 2 — Asalka Mowduuca (Background)", 1)
    bullets(doc, [
        "Dembiyada maalin walba way dhacaan — warbixinnadooda waxay ku yimaadaan qoraal (social media, email, warbaahin).",
        "NLP (Natural Language Processing) wuxuu kombiyuutarka u suurtageliyaa inuu fahmo luqadda dadka.",
        "Machine Learning wuxuu bartaa qaab-dhismeedka qoraallada dembiga iyo kuwa aan dembiga ahayn.",
        "Soomaaliya: qoraallo badan oo Soomaali + Ingiriisi ah — habka gacanta ma dabooli karo mugga.",
    ])
    body(doc, "Hadaf guud: In la dhiso nidaam casri ah oo ka caawiya hay'adaha amniga inay si degdeg ah u kala saaraan warbixinnada.")
    page_break(doc)

    # SLIDE 3
    heading(doc, "SLIDE 3 — Dhibaatada (Problem Statement)", 1)
    bullets(doc, [
        "Kala saarid GACANTA ah: gaabis, khaladaad (daal), qaalin (shaqaale aqoon leh).",
        "Tusaale: 500 qoraal/maalin × 2 daqiiqo = 16+ saacadood oo kaliya kala saarid.",
        "Nidaamyo jira: Ingiriisi kaliya, keyword filter fudud, ma jiro workflow baaritaan.",
        "Gap: Ma jiro platform Soomaali + ML + monitoring + cases hal meel.",
    ])
    body(doc, "Su'aasha cilmi-baarista: Sidee loo dhisi karaa nidaam NLP ah oo otomaatig ah oo ku habboon Soomaaliya?")
    page_break(doc)

    # SLIDE 4
    heading(doc, "SLIDE 4 — Ujeeddooyinka (Objectives)", 1)
    heading(doc, "Ujeeddo Guud", 2)
    body(doc, "In la dhiso framework otomaatig ah oo ML + NLP ku kala saara warbixinnada dembiga.")
    heading(doc, "Ujeeddooyin Gaar ah", 2)
    bullets(doc, [
        "In la yareeyo ku tiirsanaanta habka gacanta.",
        "In la ururiyo oo la diyaariyo xogta qoraalka dembiga.",
        "In la soo saaro features muhiim ah oo kala saarid sax ah.",
        "In la dhiso nidaam NLP ah oo otomaatig ah.",
    ])
    page_break(doc)

    # SLIDE 5
    heading(doc, "SLIDE 5 — Xalka: BAREAI", 1)
    body(doc, "BAREAI = Bare AI Crime Intelligence Platform — nidaam web ah oo dhammaystiran.")
    bullets(doc, [
        "Kala saarid: Crime vs Not-crime + confidence score.",
        "Input: Text, URL, File (PDF/DOCX), Batch.",
        "Hybrid: ML model + ereyo Soomaali + goobo + blacklist.",
        "Facebook monitoring (Puppeteer) — ogaanshaha hore.",
        "Cases, Notifications, Reports — workflow baaritaan.",
    ])
    page_break(doc)

    # SLIDE 6
    heading(doc, "SLIDE 6 — Qaab-dhismeedka Nidaamka (Architecture)", 1)
    bullets(doc, [
        "Frontend: React 19 + Vite (port 5173) — Dashboard, Analysis, Cases…",
        "Backend: Node.js + Express (port 5000) — Auth, API, Facebook monitor.",
        "AI Service: Python Flask (port 5001) — crime_model.pkl + vectorizer.pkl.",
        "Database: MongoDB — crime_detection_system.",
        "Amni: JWT + bcrypt + roles (admin, investigator, analyst).",
    ])
    body(doc, "Saddex-laab (3-tier) + microservice AI — model-ka si madaxbannaan ayaa loo cusboonaysiin karaa.")
    page_break(doc)

    # SLIDE 7
    heading(doc, "SLIDE 7 — Habka Cilmi-baarista (Methodology)", 1)
    bullets(doc, [
        "Dataset: dataset.csv.csv — qoraallo la calaamadiyay (crime / not crime).",
        "Preprocessing: nadiifin, TF-IDF vectorization, train-test split 80/20.",
        "Models: SVM, Random Forest, BERT — isbarbardhig.",
        "Metrics: Accuracy, Precision, Recall, F1, AUC-ROC.",
        "Kan ugu fiican → deploy Flask /predict endpoint.",
    ])
    page_break(doc)

    # SLIDE 8
    heading(doc, "SLIDE 8 — Astaamaha Muhiimka ah (Key Features)", 1)
    bullets(doc, [
        "Analysis Center — 4 nooc oo input ah.",
        "Dashboard — KPIs, trends, keywords.",
        "Blacklist — pages, persons, keywords, websites.",
        "Notifications — digniino degdeg ah.",
        "Case Management — investigator assignment + verdict.",
        "PDF Reports — weekly, monthly, individual.",
    ])
    body(doc, "Demo: Qoraal Soomaali ah → natiijo Crime + keywords + goob.")
    page_break(doc)

    # SLIDE 9
    heading(doc, "SLIDE 9 — Natiijooyinka & Tijaabada", 1)
    bullets(doc, [
        "Model comparison: SVM, RF, BERT (tirooyinka notebook-ka).",
        "20 test cases — dhammaan Pass.",
        "Wakhtiga jawaabta: < 3 ilbiriqsi (AI), < 30 ilbiriqsi (end-to-end).",
        "Keyword fallback haddii AI uu dhaco — nidaamku wuu sii shaqaynayaa.",
    ])
    body(doc, "BAREAI wuxuu ka dhakhso badan yahay kala saarid gacanta oo wuxuu bixiyaa isku mid (consistency).")
    page_break(doc)

    # SLIDE 10
    heading(doc, "SLIDE 10 — Gabagabo & Mustaqbal", 1)
    heading(doc, "Gabagabo", 2)
    bullets(doc, [
        "Dhibaatada kala saarid gacanta waa la xaliyay iyadoo la adeegsanayo NLP + ML.",
        "BAREAI waa decision-support — bini'aadamku go'aanka ugu dambeeya wuu qaataa.",
        "Platform dhammaystiran: ma aha kaliya model.",
    ])
    heading(doc, "Mustaqbal (Future Work)", 2)
    bullets(doc, [
        "Multi-class crime types (rabshad, xatooyo, maandooriye…).",
        "Dataset Soomaali oo ballaaran + Afri-BERTa.",
        "Mobile app + Netlify deploy.",
    ])
    centered(doc, "Mahadsanid — Su'aalo?", 14, bold=True)

    return doc


def build_script() -> Document:
    doc = Document()
    style_doc(doc)

    heading(doc, "SCRIPT AF SOOMAALI — 5 DAQIIQO (Pre-Panel)", 1)
    body(doc, "Akhriso si dabiici ah. Waqtiga qiyaasta: 4:30 – 5:30 daqiiqo. Hoos ka akhriso qayb kasta.")

    script = """
[SLIDE 1 — 15 ilbiriqsi]
Assalaamu calaykum. Magacaygu waa [magacaaga]. Waxaan idinla wadaagayaa mashruucayaga qalin-jebinta: BAREAI — Automatic Classification of Crime-Related Text Reports Using NLP. Waa nidaam casri ah oo ka caawiya hay'adaha amniga inay si otomaatig ah u kala saaraan warbixinnada qoraalka ah. Kooxdayadu waxay ka kooban tahay afar arday oo ka tirsan JUST, Faculty of Computer & Information Technology.

[SLIDE 2 — 40 ilbiriqsi]
Mowduucan wuxuu ka yimaadaa xaqiiqda nolosha maalinlaha ah. Dembiyada maalin walba way dhacaan, warbixinnadoodana waxay ku yimaadaan qoraal — Facebook, email, warbaahinta, iyo foomo online ah. NLP waa qayb ka mid ah sirdoonka macmalka ah oo kombiyuutarka u suurtagelisa inuu fahmo luqadda dadka. Machine Learning wuxuu bartaa qaabka qoraallada dembiga iyo kuwa aan dembiga ahayn. Soomaaliya gaar ahaan, qoraalladu waa isku dhafan yihiin Soomaali iyo Ingiriisi — taasina waxay ka dhigaysaa habka gacanta mid aan la qabsan karin mugga xogta ee maalin walba soo socota.

[SLIDE 3 — 50 ilbiriqsi]
Hadda aan si cad u sheegno dhibaatada. Hay'adaha amniga waxay helaan kumanaan qoraal maalin kasta, laakiin intooda badan waxaa gacanta u kala saaraa qof — inuu go'aamiyo: dembi baa la xiriira mise ma aha. Habkani wuxuu leeyahay saddex dhibaato oo waaweyn: waa gaabis, sababtoo ah waxay qaadataa waqti badan; waa khaladaad, sababtoo ah daalka iyo is-maangal la'aanta; waa qaalin, sababtoo ah shaqaalaha aqoon leh waxay ku mashquulaan kala saarid halkii ay baaritaan sameyn lahaayeen. Tusaale fudud: haddii 500 qoraal la helo maalintii, qof kasta laba daqiiqo ku qaato, waxaa loo baahan yahay in ka badan lix iyo toban saacadood oo kaliya kala saarid. Nidaamyo kale oo jira waxay inta badan ku shaqeeyaan Ingiriisi kaliya, ama keyword filter fudud, mana laha workflow baaritaan. Gap-ka aan buuxinayno waa: ma jiro platform isku dhafan oo Soomaali taageera, ML isticmaala, Facebook la monitor gareeyo, iyo cases la maamulo — hal meel.

[SLIDE 4 — 35 ilbiriqsi]
Ujeeddadeenna guud waa in la dhiso framework otomaatig ah oo ML iyo NLP ku kala saara warbixinnada dembiga. Ujeeddooyinkeenna gaarka ah waa afar: in la yareeyo ku tiirsanaanta gacanta; in la ururiyo oo la diyaariyo xogta; in la soo saaro features sax ah; iyo in la dhiso nidaam NLP ah oo otomaatig ah. Su'aalaha cilmi-baaristu waxay ku wareegayaan: sidee loo sameeyaa waxan oo dhan si wax ku ool ah?

[SLIDE 5 — 45 ilbiriqsi]
Xalkayagu waa BAREAI — Bare AI Crime Intelligence Platform. Ma aha kaliya model ML; waa platform web ah oo dhammaystiran. Wuxuu qoraalka u kala saaraa laba: Crime ama Not-crime, wuxuuna bixiyaa confidence score. Waxaad soo gudbin kartaa text toos ah, URL, file PDF ama DOCX, ama batch. Waxa aan ku darnay hybrid approach: model-ka ML waxaa lagu daray ereyo Soomaali, goobo, iyo blacklist. Waxaan sidoo kale haysannaa Facebook monitoring si aan u ogaanno dembiyada ka hor inta aan la soo gudbin. Waxaa jira cases, notifications, iyo reports PDF ah.

[SLIDE 6 — 40 ilbiriqsi]
Qaab-dhismeedka nidaamku waa saddex-laab. Frontend-ka React iyo Vite ayaa loo isticmaalay — dashboard, analysis, cases, iwm. Backend-ka Node.js iyo Express wuxuu maamulaa authentication, API, iyo Facebook monitoring. AI service-ka Python Flask wuxuu ku shaqeeyaa port 5001, wuxuuna xambaarsan yahay crime_model.pkl iyo vectorizer.pkl. Xogta waxaa lagu kaydiyaa MongoDB. Amniga waxaa lagu xaqiijiyaa JWT, bcrypt, iyo roles: admin, investigator, iyo analyst.

[SLIDE 7 — 40 ilbiriqsi]
Habka cilmi-baarista: waxaan isticmaalnay dataset la calaamadiyay — crime iyo not crime. Preprocessing waxaa ka mid ah nadiifinta qoraalka, TF-IDF vectorization, iyo train-test split siddeed iyo labaatan. Waxaan isbarbardhignay saddex model: SVM, Random Forest, iyo BERT. Waxaan qiimeynay accuracy, precision, recall, F1, iyo AUC. Model-ka ugu fiican waxaan u deploy gareynay Flask endpoint-ka /predict.

[SLIDE 8 — 35 ilbiriqsi]
Astaamaha muhiimka ah waxaa ka mid ah Analysis Center oo leh afar nooc oo input ah, Dashboard oo muujiya KPIs iyo trends, Blacklist, Notifications, Case Management, iyo PDF Reports. Demo ahaan, markaad qoraal Soomaali ah geliso oo aad riixdo Analyze, nidaamku wuxuu soo celinayaa Crime, confidence, keywords, iyo goob haddii la helo.

[SLIDE 9 — 35 ilbiriqsi]
Natiijooyinka: model-yada waa la isbarbardhigay, test cases-na waa la mariyay — labaatan test, dhammaantood Pass. Jawaabtu waxay qaadataa ka yar saddex ilbiriqsi AI-ga, iyo ka yar soddon ilbiriqsi end-to-end. Haddii AI service uu dhaco, backend wuxuu isticmaalaa keyword fallback — nidaamku ma istaagayo.

[SLIDE 10 — 30 ilbiriqsi]
Gabagabo ahaan, BAREAI wuxuu xallinayaa dhibaatada kala saarid gacanta iyadoo la adeegsanayo NLP iyo ML. Waa decision-support tool — go'aanka ugu dambeeya wuxuu weli la xiriiraa sarkaalka baaritaanka. Mustaqbalka waxaan rabnaa multi-class classification, dataset Soomaali oo ballaaran, iyo mobile app. Mahadsanid — waxaan diyaar u nahay su'aalihiina.
"""
    p = doc.add_paragraph(script.strip())
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    page_break(doc)
    heading(doc, "Su'aalaha La Diyaariyay (Quick Reference)", 1)
    qa = [
        ("Waa maxay dhibaatada?", "Kala saarid gacanta — gaabis, khaladaad, qaalin."),
        ("Maxay BAREAI ka duwan tahay keyword filter?", "ML wuxuu fahmaa macnaha; keywords waa enrichment."),
        ("Maxaa dhacaya haddii AI uu qalad saxo?", "Investigator wuu xaqiijinayaa — case + verdict."),
        ("Sidee Facebook loo monitor gareeyaa?", "Puppeteer — bogag dadweyne, ma aha farriimo gaar ah."),
        ("Waa maxay TF-IDF?", "Qoraalka u beddelka tirooyin si model-ku u barto."),
    ]
    for q, a in qa:
        heading(doc, f"Su'aal: {q}", 3)
        body(doc, f"Jawaab: {a}")

    return doc


def main() -> None:
    slides = build_slides()
    slides_path = OUT_DIR / "BAREAI_PrePanel_Presentation.docx"
    slides.save(str(slides_path))

    script = build_script()
    script_path = OUT_DIR / "BAREAI_PrePanel_Script_Somali.docx"
    script.save(str(script_path))

    # Combined single file for convenience
    combined = Document()
    style_doc(combined)
    heading(combined, "QAYBTA 1 — SLIDES", 1)
    for para in build_slides().paragraphs:
        combined.add_paragraph(para.text)
    page_break(combined)
    heading(combined, "QAYBTA 2 — SCRIPT", 1)
    for para in build_script().paragraphs:
        combined.add_paragraph(para.text)
    combined_path = OUT_DIR / "BAREAI_PrePanel_Complete.docx"
    combined.save(str(combined_path))

    print(f"Saved: {slides_path}")
    print(f"Saved: {script_path}")
    print(f"Saved: {combined_path}")


if __name__ == "__main__":
    main()
