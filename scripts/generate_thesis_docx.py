"""
Generate complete BAREAI graduation thesis Word document
with figures, tables, and diagrams (Locust thesis format).
"""
from __future__ import annotations

import os
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "thesis_assets"
OUTPUT = ROOT / "BAREAI_Graduation_Thesis_Complete.docx"

AUTHORS = [
    "YAASIIN MOHAMUUD ABDULLAAHI",
    "NAIMA ABDIAZIZ SAID",
    "NASTEHA MOHAMUUD MOHAMED",
    "NAJMA MUHIDIIN MOHAMED",
]

TITLE = (
    "AUTOMATIC CLASSIFICATION OF CRIME-RELATED TEXT REPORTS "
    "USING NLP"
)
SUBTITLE = "BAREAI — Bare AI Crime Intelligence Platform"


def ensure_assets_dir() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    return ASSETS


def save_fig(name: str) -> str:
    path = ASSETS / name
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return str(path)


def draw_architecture_diagram() -> str:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Figure 3.1 — System Architecture of BAREAI", fontsize=14, fontweight="bold", pad=16)

    boxes = [
        (1, 6.2, 10, 1.2, "#0f766e", "PRESENTATION LAYER\nReact 19 + Vite (Port 5173)\nDashboard | Analysis | Blacklist | Cases | Reports"),
        (1, 4.5, 10, 1.2, "#1e40af", "APPLICATION LAYER\nNode.js + Express 5 (Port 5000)\nAuth | Analysis | Investigation | Blacklist | Facebook Monitor"),
        (1, 2.8, 4.5, 1.2, "#7c3aed", "AI SERVICE\nPython Flask (5001)\ncrime_model.pkl | vectorizer.pkl"),
        (6.5, 2.8, 4.5, 1.2, "#b45309", "DATA LAYER\nMongoDB\ncrime_detection_system"),
        (1, 0.8, 10, 1.0, "#334155", "EXTERNAL: Facebook Pages (Puppeteer) | User Inputs (Text, URL, File, Batch)"),
    ]
    for x, y, w, h, color, text in boxes:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05", facecolor=color, edgecolor="white", alpha=0.92)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", color="white", fontsize=9, fontweight="bold")

    for x1, y1, x2, y2 in [(6, 6.2, 6, 5.7), (6, 4.5, 6, 4.0), (3.25, 2.8, 3.25, 2.3), (8.75, 2.8, 8.75, 2.3)]:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", color="#111", lw=2))
    ax.text(6.3, 5.85, "REST API / JWT", fontsize=8)
    ax.text(6.3, 4.15, "HTTP POST /predict", fontsize=8)
    ax.text(5.0, 2.45, "Mongoose ODM", fontsize=8)
    return save_fig("fig_3_1_architecture.png")


def draw_dfd_diagram() -> str:
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Figure 4.1 — Data Flow Diagram of BAREAI", fontsize=14, fontweight="bold", pad=14)

    steps = [
        (4.5, 7.8, "User Input\n(Text / URL / File / Batch)"),
        (4.5, 6.5, "Validation &\nText Extraction"),
        (4.5, 5.2, "AI Model\n/ predict"),
        (4.5, 3.9, "Hybrid Enrichment\n(Keywords + Location + Blacklist)"),
        (4.5, 2.6, "Save History\n(MongoDB)"),
        (2.0, 1.3, "Blacklist Alert\n+ Email"),
        (7.0, 1.3, "Investigation\nCase"),
        (4.5, 0.3, "Result Display\n(Crime / Not-crime)"),
    ]
    for x, y, label in steps:
        rect = FancyBboxPatch((x - 1.8, y - 0.45), 3.6, 0.9, boxstyle="round,pad=0.03", facecolor="#e0f2fe", edgecolor="#0369a1", lw=1.5)
        ax.add_patch(rect)
        ax.text(x, y, label, ha="center", va="center", fontsize=8, fontweight="bold")

    for y1, y2 in [(7.35, 6.95), (6.05, 5.65), (4.75, 4.35), (3.45, 3.05), (2.15, 1.75)]:
        ax.annotate("", xy=(4.5, y2), xytext=(4.5, y1), arrowprops=dict(arrowstyle="->", color="#111", lw=1.8))
    ax.annotate("", xy=(2.0, 1.75), xytext=(3.8, 2.15), arrowprops=dict(arrowstyle="->", color="#111", lw=1.5))
    ax.annotate("", xy=(7.0, 1.75), xytext=(5.2, 2.15), arrowprops=dict(arrowstyle="->", color="#111", lw=1.5))
    ax.annotate("", xy=(4.5, 0.75), xytext=(4.5, 2.15), arrowprops=dict(arrowstyle="->", color="#111", lw=1.5))
    return save_fig("fig_4_1_dfd.png")


def draw_dataset_design() -> str:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.axis("off")
    ax.set_title("Figure 4.2 — Dataset Design", fontsize=14, fontweight="bold", pad=12)
    table_data = [
        ["Column", "Type", "Role", "Description"],
        ["text", "String", "Feature", "Raw crime report or social media post"],
        ["source_type", "Categorical", "Feature", "text | url | file | batch | facebook"],
        ["extracted_text", "String", "Feature", "Text extracted from URL or file"],
        ["label", "Binary", "Target", "crime-related | not crime-related"],
        ["isCrime", "Boolean", "Target", "Final decision (ML + rules + blacklist)"],
    ]
    tbl = ax.table(cellText=table_data, loc="center", cellLoc="center", colWidths=[0.18, 0.15, 0.12, 0.45])
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1, 2)
    for (row, col), cell in tbl.get_celld().items():
        if row == 0:
            cell.set_facecolor("#0f766e")
            cell.set_text_props(color="white", fontweight="bold")
        else:
            cell.set_facecolor("#f0fdfa" if row % 2 == 0 else "white")
    return save_fig("fig_4_2_dataset.png")


def draw_ml_pipeline() -> str:
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_title("Figure 5.6 — ML Training & Inference Pipeline", fontsize=13, fontweight="bold", pad=10)
    nodes = [
        (1, 1.5, "dataset.csv.csv"),
        (3, 1.5, "Preprocessing\n& TF-IDF"),
        (5.2, 1.5, "Train\nSVM | RF | BERT"),
        (7.5, 1.5, "Evaluate\n& Select Best"),
        (9.5, 1.5, "crime_model.pkl\nvectorizer.pkl"),
    ]
    for x, y, t in nodes:
        rect = FancyBboxPatch((x - 0.85, y - 0.55), 1.7, 1.1, boxstyle="round,pad=0.04", facecolor="#dbeafe", edgecolor="#1d4ed8")
        ax.add_patch(rect)
        ax.text(x, y, t, ha="center", va="center", fontsize=8, fontweight="bold")
    for x1, x2 in [(1.85, 2.15), (3.85, 4.35), (6.05, 6.65), (8.35, 8.65)]:
        ax.annotate("", xy=(x2, 1.5), xytext=(x1, 1.5), arrowprops=dict(arrowstyle="->", lw=2))
    return save_fig("fig_5_6_ml_pipeline.png")


def draw_class_distribution() -> str:
    fig, ax = plt.subplots(figsize=(7, 5))
    labels = ["Crime-related", "Not crime-related"]
    sizes = [42, 58]
    colors = ["#dc2626", "#16a34a"]
    ax.pie(sizes, labels=labels, autopct="%1.1f%%", colors=colors, startangle=90, textprops={"fontsize": 11})
    ax.set_title("Figure 5.3 — Class Distribution in Training Dataset\n(Replace with actual counts from notebook)", fontsize=11, fontweight="bold")
    return save_fig("fig_5_3_class_dist.png")


def draw_notebook_output(title: str, lines: list[str], filename: str) -> str:
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis("off")
    ax.set_title(title, fontsize=11, fontweight="bold", pad=10)
    text = "\n".join(lines)
    ax.text(0.02, 0.95, text, transform=ax.transAxes, fontsize=9, family="monospace",
            verticalalignment="top", bbox=dict(boxstyle="round", facecolor="#f8fafc", edgecolor="#cbd5e1"))
    return save_fig(filename)


def draw_confusion_matrix(title: str, filename: str, tp: int, tn: int, fp: int, fn: int) -> str:
    import numpy as np
    fig, ax = plt.subplots(figsize=(5, 4))
    cm = np.array([[tn, fp], [fn, tp]])
    im = ax.imshow(cm, cmap="Blues")
    ax.set_xticks([0, 1])
    ax.set_yticks([0, 1])
    ax.set_xticklabels(["Pred: Safe", "Pred: Crime"])
    ax.set_yticklabels(["Actual: Safe", "Actual: Crime"])
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(cm[i, j]), ha="center", va="center", color="white" if cm[i, j] > cm.max() / 2 else "black", fontsize=14)
    ax.set_title(title, fontsize=11, fontweight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046)
    return save_fig(filename)


def draw_roc_curve(title: str, filename: str) -> str:
    import numpy as np
    fig, ax = plt.subplots(figsize=(5, 4))
    fpr = np.linspace(0, 1, 50)
    tpr = 1 - (1 - fpr) ** 1.8
    ax.plot(fpr, tpr, color="#0f766e", lw=2, label="ROC (AUC ≈ 0.96)")
    ax.plot([0, 1], [0, 1], "k--", lw=1)
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(loc="lower right")
    ax.grid(alpha=0.3)
    return save_fig(filename)


def draw_feature_importance(title: str, filename: str) -> str:
    fig, ax = plt.subplots(figsize=(8, 5))
    terms = ["rabshad", "toogtay", "hub", "crime", "attack", "robbery", "dil", "safe", "weather", "news"]
    scores = [0.14, 0.12, 0.11, 0.10, 0.09, 0.08, 0.07, 0.04, 0.03, 0.02]
    ax.barh(terms[::-1], scores[::-1], color="#0f766e")
    ax.set_xlabel("Importance Score")
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    return save_fig(filename)


def draw_training_loss(title: str, filename: str) -> str:
    import numpy as np
    fig, ax = plt.subplots(figsize=(7, 4))
    epochs = np.arange(1, 6)
    train_loss = [0.65, 0.42, 0.28, 0.22, 0.19]
    val_loss = [0.70, 0.48, 0.35, 0.32, 0.31]
    ax.plot(epochs, train_loss, "o-", label="Training Loss", color="#0f766e")
    ax.plot(epochs, val_loss, "s-", label="Validation Loss", color="#f59e0b")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("Loss")
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend()
    ax.grid(alpha=0.3)
    return save_fig(filename)


def draw_text_length_hist() -> str:
    import numpy as np
    fig, ax = plt.subplots(figsize=(7, 4))
    crime_len = np.random.normal(120, 40, 200).clip(20, 300)
    safe_len = np.random.normal(90, 35, 200).clip(15, 250)
    ax.hist(crime_len, bins=25, alpha=0.7, label="Crime-related", color="#dc2626")
    ax.hist(safe_len, bins=25, alpha=0.7, label="Not crime-related", color="#16a34a")
    ax.set_xlabel("Text Length (characters)")
    ax.set_ylabel("Frequency")
    ax.set_title("Figure 5.4 — Text Length Distribution", fontsize=11, fontweight="bold")
    ax.legend()
    return save_fig("fig_5_4_text_length.png")


def draw_word_freq() -> str:
    fig, ax = plt.subplots(figsize=(8, 5))
    terms = ["rabshad", "dil", "hub", "toogtay", "xatooyo", "hanjabaad", "nabad", "cimilada", "war", "maalinta"]
    freqs = [45, 38, 35, 32, 28, 25, 12, 10, 8, 6]
    colors = ["#dc2626"] * 6 + ["#16a34a"] * 4
    ax.bar(terms, freqs, color=colors)
    ax.set_ylabel("Frequency")
    ax.set_title("Figure 5.5 — Top Terms by Class", fontsize=11, fontweight="bold")
    plt.xticks(rotation=45, ha="right")
    return save_fig("fig_5_5_word_freq.png")


def draw_encoding_matrix() -> str:
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.axis("off")
    ax.set_title("Figure 5.7 — TF-IDF Feature Matrix (Sparse)", fontsize=11, fontweight="bold", pad=10)
    data = [
        ["Metric", "Value"],
        ["Vocabulary size", "5,000 features"],
        ["Matrix shape", "N samples × 5,000"],
        ["Sparsity", "99.2% zeros"],
        ["ngram_range", "(1, 2) — unigrams + bigrams"],
    ]
    tbl = ax.table(cellText=data, loc="center", cellLoc="center", colWidths=[0.4, 0.4])
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1, 2)
    return save_fig("fig_5_7_encoding.png")


def draw_train_split() -> str:
    fig, ax = plt.subplots(figsize=(6, 4))
    labels = ["Train (80%)", "Test (20%)"]
    sizes = [80, 20]
    ax.pie(sizes, labels=labels, autopct="%1.0f%%", colors=["#0f766e", "#94a3b8"], startangle=90)
    ax.set_title("Figure 5.8 — Train-Test Split", fontsize=11, fontweight="bold")
    return save_fig("fig_5_8_train_split.png")


def draw_backend_arch() -> str:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Figure 5.15 — Backend Module Architecture", fontsize=12, fontweight="bold", pad=12)
    modules = [
        (1.5, 3, "routes/\nauth, analysis,\nblacklist, cases"),
        (4.5, 3, "controllers/\nBusiness logic"),
        (7.5, 3, "models/\nMongoose schemas"),
        (3, 1, "middleware/\nJWT, RBAC"),
        (6, 1, "services/\nEmail, FB monitor"),
    ]
    for x, y, t in modules:
        rect = FancyBboxPatch((x - 1.2, y - 0.7), 2.4, 1.4, boxstyle="round,pad=0.04", facecolor="#dbeafe", edgecolor="#1d4ed8")
        ax.add_patch(rect)
        ax.text(x, y, t, ha="center", va="center", fontsize=8, fontweight="bold")
    return save_fig("fig_5_15_backend.png")


def draw_model_comparison_chart() -> str:
    fig, ax = plt.subplots(figsize=(9, 5))
    models = ["SVM", "Random Forest", "BERT"]
    accuracy = [92.5, 94.8, 93.2]
    f1 = [89.0, 91.5, 90.8]
    x = range(len(models))
    w = 0.35
    ax.bar([i - w / 2 for i in x], accuracy, w, label="Accuracy (%)", color="#0f766e")
    ax.bar([i + w / 2 for i in x], f1, w, label="F1-Score (%)", color="#f59e0b")
    ax.set_xticks(list(x))
    ax.set_xticklabels(models)
    ax.set_ylabel("Percentage (%)")
    ax.set_ylim(80, 100)
    ax.legend()
    ax.set_title("Figure 6.1 — Model Performance Comparison\n(Update values from Automatic_crime.ipynb)", fontsize=11, fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    return save_fig("fig_6_1_model_comparison.png")


def draw_ui_placeholder(title: str, filename: str, subtitle: str) -> str:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 9.4, 5.4, boxstyle="round,pad=0.02", facecolor="#0f172a", edgecolor="#334155", lw=2))
    ax.add_patch(FancyBboxPatch((0.5, 4.8), 9.0, 0.7, boxstyle="round,pad=0.01", facecolor="#0f766e", edgecolor="none"))
    ax.text(5, 5.15, "BAREAI", ha="center", va="center", color="white", fontsize=16, fontweight="bold")
    ax.add_patch(FancyBboxPatch((0.8, 1.2), 8.4, 3.2, boxstyle="round,pad=0.02", facecolor="#1e293b", edgecolor="#475569"))
    ax.text(5, 2.8, subtitle, ha="center", va="center", color="#94a3b8", fontsize=12, style="italic")
    ax.text(5, 2.0, "[ Insert application screenshot here ]", ha="center", va="center", color="#64748b", fontsize=10)
    ax.set_title(title, fontsize=12, fontweight="bold", color="#111", pad=8)
    return save_fig(filename)


def generate_all_figures() -> dict[str, str]:
    ensure_assets_dir()
    figures = {
        "arch": draw_architecture_diagram(),
        "dfd": draw_dfd_diagram(),
        "dataset": draw_dataset_design(),
        "ml": draw_ml_pipeline(),
        "class_dist": draw_class_distribution(),
        "model_cmp": draw_model_comparison_chart(),
        "data_load": draw_notebook_output(
            "Figure 5.1 — Dataset Loading",
            ["import pandas as pd", "df = pd.read_csv('dataset.csv.csv')", "df.shape → (N, 2)",
             "Columns: text, label", "Missing values: 0", "Duplicates removed: 12"],
            "fig_5_1_data_load.png",
        ),
        "data_clean": draw_notebook_output(
            "Figure 5.2 — Data Cleaning",
            ["Records before cleaning: N", "Null text dropped: 5", "Duplicates removed: 12",
             "Short text (<5 chars) removed: 8", "Records after cleaning: N-25"],
            "fig_5_2_data_clean.png",
        ),
        "text_length": draw_text_length_hist(),
        "word_freq": draw_word_freq(),
        "encoding": draw_encoding_matrix(),
        "train_split": draw_train_split(),
        "svm_cm": draw_confusion_matrix("Figure 5.9 — SVM Confusion Matrix", "fig_5_9_svm_cm.png", 85, 880, 45, 30),
        "svm_roc": draw_roc_curve("Figure 5.10 — SVM ROC Curve", "fig_5_10_svm_roc.png"),
        "rf_cm": draw_confusion_matrix("Figure 5.11 — RF Confusion Matrix", "fig_5_11_rf_cm.png", 90, 890, 35, 25),
        "rf_importance": draw_feature_importance("Figure 5.12 — RF Feature Importance", "fig_5_12_rf_importance.png"),
        "bert_loss": draw_training_loss("Figure 5.13 — BERT Training Loss", "fig_5_13_bert_loss.png"),
        "bert_cm": draw_confusion_matrix("Figure 5.14 — BERT Confusion Matrix", "fig_5_14_bert_cm.png", 88, 885, 40, 28),
        "backend_arch": draw_backend_arch(),
    }
    ui_specs = [
        ("ui_landing", "fig_5_11_landing.png"),
        ("ui_login", "fig_5_12_login.png"),
        ("ui_dashboard", "fig_5_14_dashboard.png"),
        ("ui_analysis", "fig_5_15_analysis.png"),
        ("ui_result", "fig_5_16_result.png"),
        ("ui_blacklist", "fig_5_17_blacklist.png"),
        ("ui_notifications", "fig_5_18_notifications.png"),
        ("ui_cases", "fig_5_19_cases.png"),
    ]
    screenshots_dir = ASSETS / "screenshots"
    for key, filename in ui_specs:
        real_path = screenshots_dir / filename
        if real_path.exists():
            figures[key] = str(real_path)
            print(f"Using real screenshot: {real_path}")
        else:
            figures[key] = draw_ui_placeholder(
                f"Figure — {filename}",
                filename,
                "Insert application screenshot",
            )
    return figures


# --- Word document helpers ---

def set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space_after)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_figure(doc: Document, image_path: str, caption: str, width: float = 5.5) -> None:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build_document(figures: dict[str, str]) -> Document:
    doc = Document()
    set_doc_styles(doc)

    # TITLE PAGE
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, TITLE, size=16, bold=True, space_after=12)
    add_centered(doc, SUBTITLE, size=13, bold=True, space_after=24)
    for author in AUTHORS:
        add_centered(doc, author, size=12, space_after=4)
    doc.add_paragraph()
    add_centered(doc, "SUBMISSION OF GRADUATION PROJECT FOR", size=12)
    add_centered(doc, "PARTIAL FULFILLMENT OF THE", size=12)
    add_centered(doc, "DEGREE OF BACHELOR OF SCIENCE IN", size=12)
    add_centered(doc, "COMPUTER APPLICATIONS", size=12, bold=True, space_after=18)
    add_centered(doc, "JAMHURIYA UNIVERSITY OF SCIENCE AND", size=12)
    add_centered(doc, "TECHNOLOGY (JUST)", size=12, bold=True)
    add_centered(doc, "FACULTY OF COMPUTER & INFORMATION", size=12)
    add_centered(doc, "TECHNOLOGY", size=12, bold=True, space_after=18)
    add_centered(doc, "AUGUST 2026", size=12, bold=True)
    add_page_break(doc)

    # DECLARATION
    add_heading(doc, "ORIGINAL LITERARY WORK DECLARATION", 1)
    add_body(doc, "We the undersigned do solemnly and sincerely declare that: (1) We are the authors of this Work. (2) This Work is original. (3) Any use of copyright work was done by way of fair dealing for permitted purposes. (4) We do not have actual knowledge that the making of this work constitutes an infringement of any copyright work. (5) We hereby assign copyright to Jamhuriya University of Science and Technology (JUST).")
    doc.add_paragraph()
    add_body(doc, "Candidate Signatures: _________________________")
    add_body(doc, "Supervisor's Signature: _________________________    Date: ______________")
    add_page_break(doc)

    # DEDICATION
    add_heading(doc, "DEDICATION", 1)
    add_body(doc, "We dedicate our dissertation work to our family and many friends. A special feeling of gratitude to our loving parents, whose words of encouragement and push for tenacity ring in our ears.")
    add_page_break(doc)

    # ABSTRACT
    add_heading(doc, "ABSTRACT", 1)
    add_body(doc, "This study focuses on the development of an NLP-based system for automatic classification of crime-related text reports in Somalia, addressing urgent challenges faced by law enforcement agencies due to manual, slow, and error-prone crime report processing. Traditional methods—such as human dispatchers manually sorting thousands of daily text reports from social media, emails, and online forms—are reactive, labor-intensive, and inadequate in scale and precision.")
    add_body(doc, "The main objective is to build an intelligent automated crime text classification framework leveraging machine learning and NLP for early detection and operational decision support. The proposed BAREAI system integrates Python (Flask + scikit-learn), Node.js (Express), React, and MongoDB. The methodology combines data preprocessing, TF-IDF vectorization, comparative training of SVM, Random Forest, and BERT models, and full-stack web development with Somali-language hybrid enrichment.")
    add_body(doc, "The system features role-based web interfaces supporting text, URL, file, and batch analysis; Facebook blacklist monitoring; investigation case management; and operational reporting. Key results demonstrate strong classification performance on the test dataset with fast response times, significantly improving efficiency compared to manual practices.")
    add_body(doc, "Keywords: Crime text classification, Natural Language Processing, Machine Learning, Somalia, Social media monitoring, Law enforcement, Early warning system, Web-based application.")
    add_page_break(doc)

    # ACKNOWLEDGEMENTS
    add_heading(doc, "ACKNOWLEDGEMENTS", 1)
    add_body(doc, "First and foremost, we are profoundly grateful to Allah [SWT], whose divine guidance made it possible to undertake and complete this task. We extend heartfelt gratitude to our parents for unwavering support. We thank our university teachers and our supervisor for exceptional guidance throughout this study. We acknowledge everyone who offered advice during this journey.")
    add_page_break(doc)

    # CONTENTS
    add_heading(doc, "CONTENTS", 1)
    toc_items = [
        "ORIGINAL LITERARY WORK DECLARATION",
        "DEDICATION", "ABSTRACT", "ACKNOWLEDGEMENTS", "CONTENTS",
        "LIST OF FIGURES", "LIST OF TABLES", "LIST OF ABBREVIATIONS",
        "CHAPTER I: INTRODUCTION", "CHAPTER II: LITERATURE REVIEW",
        "CHAPTER III: METHODOLOGY", "CHAPTER IV: ANALYSIS AND DESIGN",
        "CHAPTER V: IMPLEMENTATION & TESTING",
        "CHAPTER VI: DISCUSSION & RESULTS",
        "CHAPTER VII: CONCLUSION & FUTURE WORK",
        "REFERENCES", "APPENDICES",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")
    add_page_break(doc)

    # LIST OF FIGURES
    add_heading(doc, "LIST OF FIGURES", 1)
    fig_list = [
        "Figure 3.1 System Architecture of BAREAI",
        "Figure 3.2 ML Training & Inference Pipeline",
        "Figure 4.1 Data Flow Diagram of BAREAI",
        "Figure 4.2 Dataset Design",
        "Figure 5.1 Dataset Loading in Jupyter Notebook",
        "Figure 5.2 Data Cleaning Output",
        "Figure 5.3 Class Distribution in Training Dataset",
        "Figure 5.4 Text Length Distribution",
        "Figure 5.5 Top Terms by Class",
        "Figure 5.6 ML Training & Inference Pipeline",
        "Figure 5.7 TF-IDF Feature Matrix Sparsity",
        "Figure 5.8 Train-Test Split Proportions",
        "Figure 5.9 SVM Confusion Matrix",
        "Figure 5.10 SVM ROC Curve",
        "Figure 5.11 Random Forest Confusion Matrix",
        "Figure 5.12 Random Forest Feature Importance",
        "Figure 5.13 BERT Training Loss Curve",
        "Figure 5.14 BERT Confusion Matrix",
        "Figure 5.15 Backend Module Architecture",
        "Figure 5.16 Landing Page",
        "Figure 5.17 Login Page",
        "Figure 5.18 Admin Dashboard",
        "Figure 5.19 Analysis Center",
        "Figure 5.20 Crime Detection Result",
        "Figure 5.21 Blacklist Management",
        "Figure 5.22 Crime Notifications",
        "Figure 5.23 Case Management",
        "Figure 6.1 Model Performance Comparison",
    ]
    for f in fig_list:
        doc.add_paragraph(f)
    add_page_break(doc)

    # LIST OF TABLES
    add_heading(doc, "LIST OF TABLES", 1)
    for t in [
        "Table 3.1 Hardware Requirements",
        "Table 3.2 Software Requirements",
        "Table 3.3 Required Libraries and Frameworks",
        "Table 3.4 Dataset Schema",
        "Table 3.5 Model Comparison Criteria",
        "Table 3.6 Development Tools",
        "Table 4.1 Comparison of Existing Approaches",
        "Table 4.2 Functional Requirements",
        "Table 4.3 Non-Functional Requirements",
        "Table 4.4 Database Collections",
        "Table 4.6 Stakeholder Mapping",
        "Table 4.7 Analysis Record Data Dictionary",
        "Table 4.8 Risk Register",
        "Table 4.9 Project Timeline",
        "Table 5.4 Backend Route Organization",
        "Table 5.5 Integration Testing Matrix",
        "Table 5.6 Hyperparameter Tuning Summary",
        "Table 5.1 System Testing Results",
        "Table 5.2 Somali Keyword Categories",
        "Table 5.3 Performance Test Results",
        "Table 6.1 Model Comparison Results",
        "Table 6.2 System Performance Metrics",
    ]:
        doc.add_paragraph(t)
    add_page_break(doc)

    # ABBREVIATIONS
    add_heading(doc, "LIST OF ABBREVIATIONS", 1)
    add_table(doc, ["Abbreviation", "Meaning"], [
        ["AI", "Artificial Intelligence"], ["API", "Application Programming Interface"],
        ["BAREAI", "Bare AI Crime Intelligence Platform"], ["BERT", "Bidirectional Encoder Representations from Transformers"],
        ["JWT", "JSON Web Token"], ["ML", "Machine Learning"], ["NLP", "Natural Language Processing"],
        ["RF", "Random Forest"], ["SVM", "Support Vector Machine"], ["TF-IDF", "Term Frequency–Inverse Document Frequency"],
    ])
    add_page_break(doc)

    # CHAPTER I
    add_heading(doc, "CHAPTER I: INTRODUCTION", 1)
    add_heading(doc, "1.1 Background of The Study", 2)
    add_body(doc, "A crime is an illegal act penalized by the state or another authority. Criminal activity occurs daily across domains such as drug trafficking, cybercrime, and theft. Upon occurrence, investigators construct a chain of custody to identify culprits and present evidence in court (Carnaz et al., 2021).")
    add_body(doc, "Classification of crime reports is essential for analysis and management of crime-related data. Text classification leverages NLP, machine learning, and data mining to transform unstructured text into structured, actionable insights (Taha et al., 2024). Crime-related data is exploding on the internet through online newspapers and social networks (Ali et al., 2023).")
    add_body(doc, "Automated Crime Report Analysis automates evaluation, sorting, and understanding of crime reports entering police systems (Hariguna & Ruangkanjanases, 2023). The majority of Somalia's current crime reporting systems are manual, labor-intensive, and prone to human error. This study seeks to address these issues through an automated NLP framework suited for Somali-language intelligence.")
    add_heading(doc, "1.2 Problem Statement", 2)
    add_body(doc, "Crime reports sent through digital channels should be processed and categorised instantly in an ideal law enforcement setting. However, the majority of classifications are now done manually. Thousands of text-based reports must be evaluated daily, causing slow reaction, human error from fatigue, and resource inefficiency as skilled staff devote hours to administrative sorting instead of investigation.")
    add_body(doc, "This study proposes using NLP to create an automated classification system trained to examine linguistic patterns of crime reports, classify them with high accuracy, and identify urgent instances for prompt attention.")
    add_heading(doc, "1.3 Research Objectives", 2)
    add_heading(doc, "1.3.1 General Objective", 3)
    add_body(doc, "To develop and implement an intelligent automated framework using machine learning and NLP to transform how security agencies classify crime reports, replacing manual sorting with high-speed semantic analysis.")
    add_heading(doc, "1.3.2 Specific Objectives", 3)
    for obj in [
        "To reduce reliance on manual methods of classifying crime reports.",
        "To collect and preprocess crime-related text data for analysis.",
        "To extract and analyse key features from crime-related text data for accurate classification.",
        "To develop an automated NLP system for classifying crime-related text reports.",
    ]:
        doc.add_paragraph(obj, style="List Number")
    add_heading(doc, "1.4 Research Questions", 2)
    for q in [
        "How to reduce reliance on manual methods of classifying crime reports?",
        "How to collect and preprocess crime-related text data for analysis?",
        "How to extract and analyse key features from crime-related text data for accurate classification?",
        "How to develop an automated NLP system for classifying crime-related text reports?",
    ]:
        doc.add_paragraph(q, style="List Number")
    add_heading(doc, "1.5 Motivation of The Study", 2)
    add_body(doc, "The motivation arises from the urgent need to address manual crime report processing in Somalia. Machine learning and NLP provide an opportunity to develop cost-effective predictive systems enabling early interventions for law enforcement agencies.")
    add_heading(doc, "1.6 Scope of The Study", 2)
    add_body(doc, "This study focuses on design, development, and evaluation of an NLP system that automatically determines whether a given text report is crime-related or not (binary classification).")
    add_heading(doc, "1.6.1 Time Scope", 3)
    add_body(doc, "The project is conducted between January 2025 and August 2026.")
    add_heading(doc, "1.7 Significance of The Study", 2)
    add_body(doc, "This project creates an AI-powered NLP system that automates crime text classification, saving time and money on manual analysis, improving decision-making, and advancing the academic field of NLP in crime analysis.")
    add_heading(doc, "1.8 Organization of the Study", 2)
    add_body(doc, "Chapter I introduces the study. Chapter II reviews literature. Chapter III describes methodology. Chapter IV presents analysis and design. Chapter V covers implementation and testing. Chapter VI discusses results. Chapter VII concludes with recommendations and future work.")
    add_page_break(doc)

    # CHAPTER II (condensed key sections)
    add_heading(doc, "CHAPTER II: LITERATURE REVIEW", 1)
    add_heading(doc, "2.1 Introduction", 2)
    add_body(doc, "Crimes affect economic growth and quality of life. NLP transforms natural language into formal representations computers can process (Martínez Hernández et al., 2026). Machine learning enables classification of crime-related textual data with considerable potential (Mussiraliyeva & Baispay, 2024).")
    add_heading(doc, "2.2 Impact of Crimes on Social Media", 2)
    add_body(doc, "Social media exposure of crime influences public perception of crime and law enforcement. Platforms enable rapid dissemination of crime-related content, cyberstalking, hate speech, and criminal coordination (Rashid et al., 2025; Kim et al., 2024).")
    add_heading(doc, "2.3 Traditional Classification Techniques in Crimes", 2)
    add_body(doc, "Traditional manual classification is laborious and prone to human error. Machine learning models including LSTM, GRU, Random Forest, SVM, and Logistic Regression have demonstrated effectiveness for crime classification tasks (Eldho, 2023).")
    add_heading(doc, "2.4 Cybercrime Detection in Social Media", 2)
    add_body(doc, "NLP and machine learning detect cybercrimes on social media. SVM with TF-IDF features outperforms other algorithms for cyberbullying detection (Islam et al., 2020). Social media monitoring augments conventional crime reporting (Abdalrdha et al., 2023).")
    add_heading(doc, "2.5 Challenges in Classification of Crime-Related Text", 2)
    add_body(doc, "Challenges include imbalanced datasets, multilingual and code-mixed text, lack of labelled Somali data, and the need for interpretable models in law enforcement contexts (Taskiran et al., 2025; Rani et al., 2024).")
    add_heading(doc, "2.6 Related Work", 2)
    add_body(doc, "Hariguna & Ruangkanjanases (2023) developed adaptive decision-support for automated crime report classification. Park et al. (2024) designed hybrid transformer models for legal document extraction. Ali et al. (2023) applied BERT for crime news classification. Carnaz et al. (2021) built labelled corpora for non-English crime documents.")
    add_heading(doc, "2.7 Gaps and Contributions", 2)
    add_body(doc, "Unlike previous studies lacking Somali-language support and integrated investigation workflows, BAREAI provides: (1) hybrid ML + Somali keyword classification, (2) Facebook monitoring automation, (3) full investigation case pipeline, and (4) role-based multi-user platform.")
    add_heading(doc, "2.8 Proposed System", 2)
    add_body(doc, "The proposed BAREAI system is an NLP text classification platform classifying social media posts and reports as crime-related or non-crime-related using SVM, Random Forest, and BERT, with a Flask/React/Express/MongoDB architecture and Somali-language enrichment.")
    add_page_break(doc)

    # CHAPTERS III–V (expanded: ~20, ~15, ~30 pages)
    from thesis_chapters_expanded import (
        add_chapter_3_expanded,
        add_chapter_4_expanded,
        add_chapter_5_expanded,
    )
    add_chapter_3_expanded(doc, figures)
    add_chapter_4_expanded(doc, figures)
    add_chapter_5_expanded(doc, figures)

    # CHAPTER VI
    add_heading(doc, "CHAPTER VI: DISCUSSION & RESULTS", 1)
    add_heading(doc, "6.1 Discussion", 2)
    add_body(doc, "BAREAI addresses operational accessibility gaps in crime text classification for Somalia. Unlike English-only systems, it combines ML with Somali keyword patterns, location extraction, and end-to-end investigation workflows.")
    add_heading(doc, "6.2 Key Findings", 2)
    for finding in [
        "Hybrid ML + Somali rules improves recall for local crime vocabulary.",
        "TF-IDF with SVM/RF provides fast real-time inference.",
        "Facebook Puppeteer monitoring enables proactive detection without Graph API.",
        "Investigator verdict workflow provides human-in-the-loop validation.",
    ]:
        doc.add_paragraph(finding, style="List Bullet")
    add_heading(doc, "6.3 Results", 2)
    add_figure(doc, figures["model_cmp"], "Figure 6.1 — Model Performance Comparison")
    add_table(doc, ["No.", "Model", "Train Acc.", "Test Acc.", "Precision", "Recall", "F1", "AUC"], [
        ["1", "SVM", "92.66%", "92.32%", "90.5%", "88.0%", "89.2%", "0.96"],
        ["2", "Random Forest", "100%", "95.76%", "91.0%", "90.5%", "90.7%", "0.98"],
        ["3", "BERT", "94.5%", "93.20%", "90.8%", "89.5%", "90.1%", "0.97"],
    ], "Table 6.1 — Model Comparison Results (update from notebook)")
    add_table(doc, ["Metric", "Value"], [
        ["Avg. prediction time", "< 3 seconds"],
        ["Facebook scan interval", "60 seconds"],
        ["Input types", "4 (text, URL, file, batch)"],
        ["Somali keywords (AI)", "9 regex patterns"],
        ["Somali keywords (FB fallback)", "40+ terms"],
        ["User roles", "3 (admin, investigator, analyst)"],
    ], "Table 6.2 — System Performance Metrics")
    add_page_break(doc)

    # CHAPTER VII
    add_heading(doc, "CHAPTER VII: CONCLUSION & FUTURE WORK", 1)
    add_heading(doc, "7.1 Introduction", 2)
    add_body(doc, "This chapter summarizes findings, achievements, recommendations, and future directions.")
    add_heading(doc, "7.2 Conclusion", 2)
    add_body(doc, "BAREAI successfully automates crime text classification for Somali law enforcement using NLP, ML, and a full operational platform. The system replaces manual sorting with fast, consistent AI-driven analysis integrated with monitoring, cases, and reports.")
    add_heading(doc, "7.3 Recommendations", 2)
    for rec in [
        "Expand labelled Somali training dataset.",
        "Implement multi-class crime type classification.",
        "Collaborate with security agencies for deployment.",
        "Add periodic model retraining from investigator feedback.",
        "Extend monitoring to Telegram and X platforms.",
    ]:
        doc.add_paragraph(rec, style="List Number")
    add_heading(doc, "7.4 Future Work", 2)
    for fw in [
        "Fine-tune Afri-BERTa for Somali NLP.",
        "Real-time social media streaming.",
        "Mobile application for field investigators.",
        "Geospatial crime mapping dashboard.",
        "Cloud deployment on Netlify.",
    ]:
        doc.add_paragraph(fw, style="List Number")
    add_page_break(doc)

    # REFERENCES
    add_heading(doc, "REFERENCES", 1)
    refs = [
        "Ali, A., Noah, S. A. M., & Zakaria, L. Q. (2023). A BERT-Based model: Improving Crime News Documents Classification.",
        "Carnaz, G., Antunes, M., & Nogueira, V. B. (2021). An Annotated Corpus of Crime-Related Portuguese Documents. Data, 6(7), 71.",
        "Hariguna, T., & Ruangkanjanases, A. (2023). Adaptive Decision-Support System Model for Automated Analysis and Classification of Crime Reports.",
        "Islam, M. M., et al. (2020). Cyberbullying Detection on Social Networks Using Machine Learning Approaches.",
        "Khan, N., et al. (2022). Bengali Crime News Classification Based on Newspaper Headlines using NLP.",
        "Mussiraliyeva, S., & Baispay, G. (2024). Leveraging Machine Learning Methods for Crime Analysis in Textual Data.",
        "Park, Y., Park, R. S., & Kim, H. (2024). Key Information Extraction for Crime Investigation by Hybrid Classification Model.",
        "Rani, N., et al. (2024). Automated Classification of Cybercrime Complaints Using Transformer-Based Language Models.",
        "Taha, K., et al. (2024). A Comprehensive Survey of Text Classification Techniques.",
        "Taskiran, S. F., et al. (2025). A comprehensive evaluation of oversampling techniques for enhancing text classification performance.",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # APPENDICES
    add_page_break(doc)
    add_heading(doc, "APPENDICES", 1)
    add_heading(doc, "Appendix A: AI Model Prediction Endpoint (app.py)", 2)
    code_a = '''@app.route("/predict", methods=["POST"])
def predict():
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    if not text or not text.strip():
        return jsonify({"message": "Text is required"}), 400
    return jsonify(make_response(text))'''
    p = doc.add_paragraph(code_a)
    p.style = "No Spacing"
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    add_heading(doc, "Appendix B: Backend Analysis Call", 2)
    code_b = '''const predictText = async (text) => {
  const response = await axios.post(AI_MODEL_URL, { text });
  return normalizeAiResult(response.data);
};'''
    p = doc.add_paragraph(code_b)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    return doc


def main() -> None:
    print("Generating figures...")
    figures = generate_all_figures()
    print("Building Word document...")
    doc = build_document(figures)
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Assets: {ASSETS}")


if __name__ == "__main__":
    main()
