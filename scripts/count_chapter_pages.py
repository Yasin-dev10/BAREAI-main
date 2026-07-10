"""Estimate chapter page counts in thesis docx."""
from docx import Document
from pathlib import Path

doc = Document(Path(__file__).resolve().parent.parent / "BAREAI_Graduation_Thesis_Complete.docx")
current = None
counts = {
    "III": {"words": 0, "paras": 0, "figs": 0},
    "IV": {"words": 0, "paras": 0, "figs": 0},
    "V": {"words": 0, "paras": 0, "figs": 0},
}
for p in doc.paragraphs:
    t = p.text.strip()
    if "CHAPTER III:" in t:
        current = "III"
        continue
    if "CHAPTER IV:" in t:
        current = "IV"
        continue
    if "CHAPTER V:" in t:
        current = "V"
        continue
    if "CHAPTER VI:" in t:
        current = None
        continue
    if current:
        counts[current]["paras"] += 1
        counts[current]["words"] += len(t.split())
        if t.startswith("Figure"):
            counts[current]["figs"] += 1

# Count inline images (blip elements only)
current = None
img_counts = {"III": 0, "IV": 0, "V": 0}
for p in doc.paragraphs:
    t = p.text.strip()
    if "CHAPTER III:" in t:
        current = "III"
    elif "CHAPTER IV:" in t:
        current = "IV"
    elif "CHAPTER V:" in t:
        current = "V"
    elif "CHAPTER VI:" in t:
        current = None
    if current:
        for run in p.runs:
            xml = run._element.xml
            if "<a:blip " in xml or "<pic:pic" in xml:
                img_counts[current] += 1

targets = {"III": 20, "IV": 15, "V": 30}
for ch in ["III", "IV", "V"]:
    c = counts[ch]
    # ~250 words/page + ~0.5 page per figure image
    est = c["words"] / 250 + img_counts[ch] * 0.5 + 1.5
    print(
        f"Chapter {ch}: {c['words']} words, {img_counts[ch]} images, "
        f"est ~{est:.1f} pages (target {targets[ch]})"
    )
